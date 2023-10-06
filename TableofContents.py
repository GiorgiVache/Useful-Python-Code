import docx
import re
from pathlib import Path

file = open("TableofContents.txt", 'w', encoding="utf-8")
count = 0
number = 1
Emails = []
ClearEmails = []

fi = "D:\ACS Python\Abstract Journal Symposium 2023"
abstracts = Path(fi).glob('*')
for abstract in abstracts:
    doc = docx.Document(abstract)
    var = doc.paragraphs[0].text
    var1 = var.upper()
    var2 = doc.paragraphs[2].text
    var3 = re.sub('[1234*.]', '', var2)
    var3 = re.sub('  ', ' ', var3)
    var4 = doc.paragraphs[4].text
    for par in doc.paragraphs:
        if '@' in par.text:
            Emails.append(par.text)
    file.write(var1 + "\n" + var3 + "\n" + str(number) + "\n" + var4 + '\n\n')
    number = number + 1
    count = count + 1

file.write("Total: " + str(count) + '\n')
for Email in Emails:
    ClearEmail = re.sub('Email:  ', '', Email)
    ClearEmail = re.sub('E-mail:  ', '', ClearEmail)
    ClearEmail = re.sub('E-mail: ', '', ClearEmail)
    ClearEmail = re.sub('Email: ', '', ClearEmail)
    ClearEmail = re.sub('E-mail:\xa0\xa0', '', ClearEmail)
    ClearEmail = re.sub('E-mail:  ', '', ClearEmail)
    ClearEmail = re.sub('E-Mail:  ', '', ClearEmail)
    ClearEmail = re.sub('E-mail:  ', '', ClearEmail)
    ClearEmails.append(ClearEmail)

count1 = 0
for clearemail in ClearEmails:
    count1 = count1 + 1
    file.write('\n' + str(clearemail))

file.write('\n\n' + 'Total Emails: ' + str(len(ClearEmails)))
file.write('\n\n' + str(ClearEmails))
file.close()
